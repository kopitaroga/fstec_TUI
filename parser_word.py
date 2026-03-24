import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import argparse # <-- Добавили библиотеку для аргументов
from tqdm import tqdm

# --- НАСТРОЙКИ СТОЛБЦОВ ---
# COL = {
#     'url': 13, 'f_id': 12, 'desc': 18, 'vector': 9, 'severity': 3,
#     'exploit_flag': 17, 'attacks': 19, 'icvss': 39, 'h_desc': 38, 'h_val': 36,
#     'iimp_left': 35, 'iimp_res': 36, 'iat_left': 33, 'iat_val_e': 34,
#     'type_comp': 21, 'k_k1': 27, 'k_k2': 28, 'k_res': 24,
#     'l_count': 23, 'l_l1': 29, 'l_l2': 30, 'l_res': 25,
#     'p_flag': 22, 'p_p1': 31, 'p_p2': 32, 'p_res': 26,
#     'v_total': 37
# }


class Colums: 
    """
        Класс, который содержит номера столбцов с данными
    """

    numbers = {
        'url': 13, 'f_id': 12, 'desc': 18, 'vector': 9, 'severity': 3,
        'exploit_flag': 17, 'attacks': 19, 'icvss': 39, 'h_desc': 38, 'h_val': 36,
        'iimp_left': 35, 'iimp_res': 36, 'iat_left': 33, 'iat_val_e': 34,
        'type_comp': 21, 'k_k1': 27, 'k_k2': 28, 'k_res': 24,
        'l_count': 23, 'l_l1': 29, 'l_l2': 30, 'l_res': 25,
        'p_flag': 22, 'p_p1': 31, 'p_p2': 32, 'p_res': 26,
        'v_total': 37
    }
    

    def __init__ (self, df):

        print("Значения по умолчанию:")
        self.print_numbers()
        while True:
            choice = input("Хотите ли изменить значения по умолчанию? [y/n]: ")
            if choice == "y": 
                print("\nСписок всех столбцов:")
                self.print_colum_list(df)
                self.input_new_value()
                print("\nНовые значения")
                self.print_numbers()
                break
            elif choice == "n":
                break


    def print_numbers(self):
       for key in self.numbers.keys():
        print(f"{key} : {self.numbers[key]}")
    
    
    def print_colum_list(self,df):
        cols = list(df.columns)
        for i, col in enumerate(cols, 1):
            print(f"{i}. {col}") 
    

    def input_new_value(self):
        for key in self.numbers.keys():
            while True: 
                print(f"Введите значение для {key}[default = {self.numbers[key]}]: " , end="")
                choise = input().strip()
                if choise.isdigit():
                    self.numbers[key] = int(choise)
                    break
                elif  choise == "d": 
                    break

def load_data(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.csv':
        return pd.read_csv(file_path, sep=None, engine='python', encoding='utf-8-sig')
    return pd.read_excel(file_path)

def get_v(row, col_num):
    try:
        val = row.iloc[col_num - 1]
        if pd.isna(val): return "0"
        if isinstance(val, (int, float)):
            return f"{round(val, 3)}"
        return str(val).strip()
    except:
        return "0"

def get_verdict(v_str):
    try:
        v = float(v_str.replace(',', '.'))
        if v > 8.0: return "Критический", "V > 8,0"
        if v > 5.0: return "Высокий", "V > 5,0"
        if v > 2.0: return "Средний", "V > 2,0"
        return "Низкий", "V < 2,0"
    except:
        return "Не определен", "V = ?"

def add_centered_formula(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    return p

def generate(limit=None,input_file="report.cvs"):
    output_file = 'Expert_Report_Test.docx'
    output_mini = 'Expert_Report_Mini.docx'
    output_table = 'Expert_Report_Table.docx'

    
    if not os.path.exists(input_file):
        print(f"Файл {input_file} не найден!")
        return

    print("Читаю файл данных...")
    df = load_data(input_file)
    colums = Colums(df)
   
    # ПРИМЕНЯЕМ ЛИМИТ, ЕСЛИ ОН ЗАДАН
    if limit:
        print(f"⚠️ Установлен лимит: обрабатываю только первые {limit} строк.")
        df = df.head(limit)
    


    doc = Document()
    doc_mini = Document()
    doc_table = Document()
   
    # Настройки стиля
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.first_line_indent = Cm(1.25)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 


    style = doc_mini.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.first_line_indent = Cm(1.25)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    style_t = doc_table.styles['Normal']
    style_t.font.name = 'Times New Roman'
    style_t.font.size = Pt(14)

    # Tablitsa
    table_obj = doc_table.add_table(rows=1, cols=4)
    table_obj.style = 'Table Grid' 

    hdr_cells = table_obj.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'ID BDU'
    hdr_cells[2].text = 'Показатель V'
    hdr_cells[3].text = 'Критичность'

    for index, row in tqdm(df.iterrows(), total=len(df), desc="Прогресс", unit=" стр."):
        f_id = get_v(row, colums.numbers['f_id'])
        v_total_val = get_v(row, colums.numbers['v_total'])
        verdict_word, verdict_cond = get_verdict(v_total_val)
       
        expl_val = str(row.iloc[colums.numbers['exploit_flag']-1]).strip()
        expl_text = "Существует" if expl_val == "1" else "Данные уточняются"
       
        try:
            e_float = float(str(row.iloc[colums.numbers['iat_val_e']-1]).replace(',', '.'))
        except: e_float = 0.1
       
        if e_float >= 0.6: iat_desc = "Эксплуатируется в реальных атаках"
        elif e_float >= 0.3: iat_desc = "Имеются сведения о наличие средств эксплуатации уязвимости"
        else: iat_desc = "Отсутствуют сведения об эксплуатации в реальных атаках"

        net_val = str(row.iloc[colums.numbers['p_flag']-1]).strip()



        net_text = "доступно" if net_val == "1" else "недоступно"

        try:
            l_float = float(str(row.iloc[colums.numbers['l_l2']-1]).replace(',', '.'))
        except: l_float = 0.0

        if l_float >= 1.0: comp_text = "Более 70% компонентов от общего числа компонентов в информационной системе."
        elif l_float >= 0.8: comp_text = "50-70% компонентов от общего числа компонентов в информационной системе."
        elif l_float >= 0.6: comp_text = "10-50% компонентов от общего числа компонентов в информационной системе."
        else: comp_text = "Менее 10% компонентов от общего числа компонентов в информационной системе."


        
        cvss_type = "CVSS 3.1"
        if "CVSS:4.0" in get_v(row, colums.numbers['vector']): 
            cvss_type = "CVSS 4.0"
        elif "CVSS:3.0" in get_v(row, colums.numbers['vector']) or "CVSS:3.1" in get_v(row, colums.numbers['vector']):
            cvss_type = "CVSS 3.1"
        else: 
            cvss_type = "CVSS.2.0"
            
        print(cvss_type)
        
        desc_clean = get_v(row, colums.numbers['desc'])
        print(get_v(row,colums.numbers['desc']))
        url = get_v(row,colums.numbers['url']).replace(',','.')
        # Построчная сборка (Раздел 1)
        doc.add_paragraph("1. Исходные данные об уязвимости:")
        doc.add_paragraph(f"Сведения об уязвимости, содержащиеся на сайте Банка данных угроз безопасности информации ФСТЭК России {url}")
        doc.add_paragraph(f"Идентификатор: {f_id}")
        doc.add_paragraph(f"Описание уязвимости: {desc_clean}.")
        doc.add_paragraph(f"Базовый вектор уязвимости: по {cvss_type}: {get_v(row, colums.numbers['vector'])}")
        doc.add_paragraph(f"Уровень опасности уязвимости: {get_v(row, colums.numbers['severity'])} уровень опасности (базовая оценка {cvss_type} составляет {get_v(row, colums.numbers['icvss'])})")
        doc.add_paragraph(f"Наличие эксплойта: {expl_text}")
        doc.add_paragraph(f"Эксплуатация в реальных атаках: {get_v(row, colums.numbers['attacks'])}")

        # Раздел 1.1 - 1.3
       
        doc.add_paragraph(f"1.1 Исходя из уровня опасности уязвимости {f_id} по {cvss_type}, показателю Icvss присваивается следующее значение:")
        add_centered_formula(doc, f"Icvss={get_v(row, colums.numbers['icvss'])}")

        doc.add_paragraph(f"1.2 Исходя из описания уязвимости, показателю последствий воздействий, которым подвергается информационная система при эксплуатации уязвимости присваивается значение «{get_v(row, colums.numbers['h_desc'])}» (H={get_v(row, colums.numbers['h_val'])}):")
        add_centered_formula(doc, f"Iimp={get_v(row, colums.numbers['iimp_left'])}×{get_v(row, colums.numbers['h_val'])}={get_v(row, colums.numbers['iimp_res'])}")

        doc.add_paragraph(f"1.3 Исходя из сведений о наличии эксплойта и использовании в реальных атаках показателю возможности эксплуатации уязвимости нарушителями присваивается значение «{iat_desc}» (E={get_v(row, colums.numbers['iat_val_e'])}):")
        add_centered_formula(doc, f"Iat={get_v(row, colums.numbers['iat_left'])}×{get_v(row, colums.numbers['iat_val_e'])}={get_v(row, colums.numbers['iat_val_e'])}")

        # Раздел 2
        doc.add_paragraph("2. Исходные данные об информационной системе:")
        doc.add_paragraph("2.1 Сведения о компонентах информационной системы определены в органе (организации) по результатам проведенной инвентаризации.")
        doc.add_paragraph("Исходя из данных инвентаризации определено:")
        doc.add_paragraph(f"Тип компонента информационной системы, подверженного уязвимости: {get_v(row, colums.numbers['type_comp'])}")
        k_res = float( get_v(row, colums.numbers['k_k1']))  * float(get_v(row, colums.numbers['k_k2']))
        add_centered_formula(doc, f"K=k ×K={get_v(row, colums.numbers['k_k1'])}×{get_v(row, colums.numbers['k_k2'])}={k_res}")
       
        doc.add_paragraph(f"Количество уязвимых компонентов информационной системы: {comp_text}")
        l_res = float(get_v(row, colums.numbers['l_l1'])) * float(get_v(row, colums.numbers['l_l2']))
        add_centered_formula(doc, f"L=l×L={get_v(row, colums.numbers['l_l1'])}×{get_v(row, colums.numbers['l_l2'])}={l_res}")
       
        doc.add_paragraph(f"Влияние на эффективность защиты периметра информационной системы: уязвимое программное, программно-аппаратное средство {net_text} из сети «Интернет».")
        p_res = float(get_v(row, colums.numbers['p_p1'])) * float(get_v(row, colums.numbers['p_p2']))
        add_centered_formula(doc, f"P=p×P={get_v(row, colums.numbers['p_p1'])}×{get_v(row, colums.numbers['p_p2'])}={p_res}")

        doc.add_paragraph("Таким образом, определяется показатель влияния уязвимости на функционирование информационной системы Iinfr.")
        add_centered_formula(doc, f"Iinfr=k×K+l×L+p×P={k_res}+{l_res}+{p_res}={get_v(row, colums.numbers['k_res'])}")

        # Раздел 3
        doc.add_paragraph("3. Расчет уровня критичности уязвимости программных, программно-аппаратных средств в информационной системе V осуществляется с использованием данных, полученных в предыдущих пунктах, по формуле:")
        add_centered_formula(doc, f"V =Icvss×Iinfr×(Iat+Iimp)={get_v(row, colums.numbers['icvss'])}×{get_v(row, colums.numbers['k_res'])}×({get_v(row, colums.numbers['iat_val_e'])}+{get_v(row, colums.numbers['iimp_res'])}) ≈ {v_total_val}")
        doc.add_paragraph(f"Таким образом, уровню критичности указанной уязвимости для имеющейся информационной системы присваивается значение «{verdict_word}» ({verdict_cond}).")
        
        # --- МИНИ-ОТЧЕТ (ТОЛЬКО РАЗДЕЛ 3) ---

        p_mini = doc_mini.add_paragraph()
        run_mini = p_mini.add_run(f"Идентификатор: {f_id}")
        run_mini.bold = True

        add_centered_formula(doc_mini, f"V =Icvss×Iinfr×(Iat+Iimp)={get_v(row, colums.numbers['icvss'])}×{get_v(row, colums.numbers['k_res'])}×({get_v(row, colums.numbers['iat_val_e'])}+{get_v(row, colums.numbers['iimp_res'])}) ≈ {v_total_val}")
        doc_mini.add_paragraph(f"Уровень критичности указанной уязвимости для имеющейся информационной системы присваивается значение «{verdict_word}» ({verdict_cond}).")

        row_cells = table_obj.add_row().cells
        row_cells[0].text = str(index+1)
        row_cells[1].text = f_id
        row_cells[2].text = v_total_val
        row_cells[3].text = verdict_word

        if index < len(df) - 1:
            doc.add_page_break()
            #doc_mini.add_page_break()

    doc.save(output_file)
    doc_mini.save(output_mini)
    doc_table.save(output_table)
    print(f"\nОтчеты сохранены: {output_file} и {output_mini} и {output_table}")


if __name__ == "__main__":
    # Настройка парсера аргументов
    parser = argparse.ArgumentParser(description="Генератор отчетов по уязвимостям")
    parser.add_argument("input_file", help="Путь к файлу для отчета (CSV)")
    parser.add_argument("--limit", type=int, help="Ограничить количество обрабатываемых строк для теста")
   
    args = parser.parse_args()
   
    generate(limit=args.limit,input_file=args.input_file)

